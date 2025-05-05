#!/Users/peyton.hamilton/miniconda3/envs/trace/bin/python

import tkinter as tk
import tkinter.ttk as ttk
from tkinter import messagebox, filedialog
import requests
import base64
import os
import time
import glob
from datetime import datetime
import threading
import subprocess
import sys
import xlsxwriter
import certifi
import ssl
import warnings
from urllib3.exceptions import InsecureRequestWarning
from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from update_manager import UpdateManager

# TODO: Properly fix SSL certificate verification in future update
# Current workaround: Disable SSL verification with warning
# This is temporary and should be replaced with proper certificate handling
warnings.filterwarnings('always', category=InsecureRequestWarning)

def create_ssl_context():
    """Create a secure SSL context using application certificates."""
    try:
        context = ssl.create_default_context()
        # Try to use application certificates first
        app_cert_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "certs", "cacert.pem")
        if os.path.exists(app_cert_path):
            context.load_verify_locations(cafile=app_cert_path)
        else:
            # Fall back to certifi's certificates
            context.load_verify_locations(cafile=certifi.where())
        return context
    except Exception as e:
        print(f"Warning: Could not load SSL certificates: {e}")
        # Fall back to system certificates
        context = ssl.create_default_context()
        context.load_default_certs()
        return context

# Create a session with proper SSL verification
session = requests.Session()
try:
    app_cert_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "certs", "cacert.pem")
    if os.path.exists(app_cert_path):
        session.verify = app_cert_path
    else:
        session.verify = certifi.where()
except Exception as e:
    print(f"Warning: Could not set up SSL verification: {e}")
    session.verify = True  # Fall back to system certificates

#################################
#         MAIN APPLICATION      #
#################################
class TraceApp(tk.Tk):
    def __init__(self):
        print("Starting TraceApp initialization...")
        super().__init__()
        
        # Initialize the main application window
        self.setup_main_window()
        
        # Initialize the update manager
        self.update_manager = UpdateManager(
            repo_owner="phamilton09",
            repo_name="Trace",
            current_version="1.0.0"
        )
        
        # Check for updates on startup
        self.check_for_updates()
        
        # Initialize investigation directory
        self.investigation_dir = None  # Will be set when CustomerID_Name is entered
    
    def setup_main_window(self):
        """Set up the main application window."""
        self.title("Trace - Investigation Toolkit")
        self.geometry("900x600")

        # Initialize and configure style
        self.style = ttk.Style(self)
        
        # Use system theme if available, fall back to clam
        available_themes = self.style.theme_names()
        if "system" in available_themes:
            self.style.theme_use("system")
        elif "clam" in available_themes:
            self.style.theme_use("clam")
        
        # Keep track of dark mode state (default to system setting)
        self.dark_mode_enabled = False
        try:
            if sys.platform == "darwin":
                # Check macOS dark mode
                result = subprocess.run(['defaults', 'read', '-g', 'AppleInterfaceStyle'], 
                                     capture_output=True, text=True)
                self.dark_mode_enabled = result.stdout.strip() == "Dark"
            # Add Windows dark mode detection here if needed
        except:
            pass  # Fall back to light mode if can't detect

        # Initialize base colors
        self.style.configure(".",
            background="#ffffff",
            foreground="#000000",
            fieldbackground="#ffffff",
            selectbackground="#0078d7",
            selectforeground="#ffffff"
        )

        # Apply theme before creating widgets
        self.apply_dark_mode(self.dark_mode_enabled)
        self.setup_custom_styles()

        # -- Top Frame for CustomerID_Name and Run Case
        top_frame = ttk.Frame(self)
        top_frame.pack(side="top", fill="x", pady=5)

        # Add CustomerID_Name field
        ttk.Label(top_frame, text="CustomerID_Name:").pack(side="left", padx=(10, 5))
        self.customer_id_name = ttk.Entry(top_frame, width=30)
        self.customer_id_name.pack(side="left", padx=(0, 20))
        
        # Add placeholder text for CustomerID_Name
        self.customer_id_name.bind('<FocusIn>', lambda e: self.show_placeholder(self.customer_id_name, "Enter CustomerID_Name"))
        self.customer_id_name.bind('<FocusOut>', lambda e: self.hide_placeholder(self.customer_id_name, "Enter CustomerID_Name"))
        self.customer_id_name.bind('<KeyRelease>', lambda e: self.setup_investigation_dir())  # Update directory when CustomerID_Name changes
        self.hide_placeholder(self.customer_id_name, "Enter CustomerID_Name")

        # Add label to explain Run Case functionality
        ttk.Label(top_frame, text="Double-click tabs to select for 'Run Case'").pack(side="left", padx=10)
        self.run_case_button = ttk.Button(top_frame, text="Run Case", command=self.run_selected_tabs)
        self.run_case_button.pack(side="right", padx=10)

        # -- Notebook container
        self.container = ttk.Notebook(self)
        self.container.pack(fill="both", expand=True)

        # Create frames/pages
        self.screenshot_pdf_frame = ScreenshotToPDFFrame(self.container, self)
        self.customer_summary_frame = CustomerSummaryFrame(self.container, self)
        self.alert_frame = AlertFrame(self.container, self)
        self.transaction_frame = TransactionCSVFrame(self.container, self)

        # Add frames as tabs with initial text
        self.tab_texts = [
            "1) Screenshots to PDF",
            "2) Research",
            "3) Alert Templates",
            "4) Transaction CSV"
        ]
        
        # Add tabs with minimum width to prevent text cutoff
        self.container.add(self.screenshot_pdf_frame, text=self.tab_texts[0], padding=10)
        self.container.add(self.customer_summary_frame, text=self.tab_texts[1], padding=10)
        self.container.add(self.alert_frame, text=self.tab_texts[2], padding=10)
        self.container.add(self.transaction_frame, text=self.tab_texts[3], padding=10)

        # Keep references to frames in a list (in the same order as the tabs)
        self.frames = [
            self.screenshot_pdf_frame,
            self.customer_summary_frame,
            self.alert_frame,
            self.transaction_frame
        ]

        # Track which tabs are selected for "Run Case"
        self.selected_tabs = set()

        # Bind a double-click event on the Notebook's tab area
        self.container.bind("<Double-1>", self.on_tab_double_click)

        # Add a simple menu with a dark mode toggle
        menubar = tk.Menu(self, tearoff=False)
        
        view_menu = tk.Menu(menubar, tearoff=False)
        view_menu.add_command(label="Toggle Dark Mode", command=self.toggle_dark_mode)
        menubar.add_cascade(label="View", menu=view_menu)

        help_menu = tk.Menu(menubar, tearoff=False)
        help_menu.add_command(label="About Trace", command=self.show_about)
        menubar.add_cascade(label="Help", menu=help_menu)
        
        self.config(menu=menubar)

        # Add custom style for placeholder frames
        self.style.configure("Placeholder.TFrame", background="#f0f0f0")

    def get_file_prefix(self):
        """Get the current CustomerID_Name prefix for file naming."""
        prefix = self.customer_id_name.get().strip()
        if not prefix:
            return ""
        return f"{prefix}_"
    
    def get_investigation_dir(self):
        """Get the investigation directory path based on CustomerID_Name."""
        prefix = self.customer_id_name.get().strip()
        if not prefix:
            return os.path.expanduser("~/Desktop/Investigation_File")
        return os.path.expanduser(f"~/Desktop/{prefix}")
    
    def setup_investigation_dir(self):
        """Set up the investigation directory when CustomerID_Name changes."""
        self.investigation_dir = self.get_investigation_dir()
        # Remove automatic directory creation
        # os.makedirs(self.investigation_dir, exist_ok=True)
    
    def show_placeholder(self, entry, placeholder):
        """Show placeholder text in an entry field."""
        if entry.get() == placeholder:
            entry.delete(0, tk.END)
            # Set text color to contrast with background
            bg_color = entry.cget('background')
            text_color = '#000000' if bg_color == '#ffffff' else '#ffffff'
            entry.config(foreground=text_color)
    
    def hide_placeholder(self, entry, placeholder):
        """Hide placeholder text in an entry field."""
        if not entry.get():
            entry.insert(0, placeholder)
            entry.config(foreground='gray')
            # Set text color to contrast with background when typing
            entry.bind('<KeyPress>', lambda e: self.set_contrasting_text_color(entry))
    
    def set_contrasting_text_color(self, entry):
        """Set text color to contrast with the background."""
        bg_color = entry.cget('background')
        text_color = '#000000' if bg_color == '#ffffff' else '#ffffff'
        entry.config(foreground=text_color)

    ############################
    #       TAB SELECTION      #
    ############################
    def on_tab_double_click(self, event):
        """Toggle which tabs are selected for 'Run Case' by double-click."""
        x, y = event.x, event.y
        tab_index = self.container.index(f"@{x},{y}")
        if tab_index == -1:
            return  # Click was not on a valid tab

        if tab_index in self.selected_tabs:
            # Already selected, so unselect
            self.selected_tabs.remove(tab_index)
            self.set_tab_color(tab_index, selected=False)
        else:
            # Not selected, so select
            self.selected_tabs.add(tab_index)
            self.set_tab_color(tab_index, selected=True)

    def set_tab_color(self, tab_index, selected):
        """
        Change the text of a tab to indicate whether it's selected for "Run Case."
        """
        try:
            if selected:
                # Add a star to indicate selection with subtle highlight
                self.container.tab(tab_index, text=f"★ {self.tab_texts[tab_index]}")
                self.style.configure("TNotebook.Tab", 
                                   background="#2a2a2a",  # Slightly lighter for selected tab
                                   foreground="#000000")  # Black text for selected tab
            else:
                # Remove the star and revert to default style
                self.container.tab(tab_index, text=self.tab_texts[tab_index])
                self.style.configure("TNotebook.Tab", 
                                   background="#1a1a1a",  # Dark background for tab
                                   foreground="#000000")  # Black text for tab
        except Exception as e:
            print(f"Error updating tab text: {e}")  # Debug output if needed

    def run_selected_tabs(self):
        """Runs the 'run()' method on all frames corresponding to selected tabs."""
        if not self.selected_tabs:
            messagebox.showinfo("No Tabs Selected", "Please double-click one or more tabs before using 'Run Case'.")
            return

        for tab_index in self.selected_tabs:
            try:
                frame = self.frames[tab_index]
                frame.run()
            except Exception as e:
                messagebox.showerror("Error Running Tab", f"Error running tab {tab_index}:\n{e}")

    ############################
    #       DARK MODE ETC.     #
    ############################
    def toggle_dark_mode(self):
        """Toggle dark mode on/off."""
        self.dark_mode_enabled = not self.dark_mode_enabled
        self.apply_dark_mode(self.dark_mode_enabled)

    def apply_dark_mode(self, enable_dark):
        """Apply or remove dark theme colors."""
        # Define colors based on theme
        colors = {
            'dark': {
                'bg': "#2b2b2b",
                'fg': "#ffffff",
                'entry_bg': "#404040",
                'entry_fg': "#ffffff",
                'select_bg': "#505050",
                'select_fg': "#ffffff",
                'tab_bg': "#1a1a1a",
                'tab_fg': "#ffffff",
                'disabled_bg': "#2b2b2b",
                'disabled_fg': "#808080"
            },
            'light': {
                'bg': "#ffffff",
                'fg': "#000000",
                'entry_bg': "#ffffff",
                'entry_fg': "#000000",
                'select_bg': "#0078d7",
                'select_fg': "#ffffff",
                'tab_bg': "#f0f0f0",
                'tab_fg': "#000000",
                'disabled_bg': "#f0f0f0",
                'disabled_fg': "#808080"
            }
        }
        
        theme = colors['dark'] if enable_dark else colors['light']
        
        # Configure base styles
        self.style.configure(".",
            background=theme['bg'],
            foreground=theme['fg']
        )
        
        # Configure specific widget styles
        self.style.configure("TLabel",
            background=theme['bg'],
            foreground=theme['fg']
        )
        
        self.style.configure("TButton",
            background=theme['bg'],
            foreground=theme['fg']
        )
        
        self.style.configure("TFrame",
            background=theme['bg']
        )
        
        self.style.configure("TLabelframe",
            background=theme['bg']
        )
        
        self.style.configure("TLabelframe.Label",
            background=theme['bg'],
            foreground=theme['fg']
        )
        
        # Configure Entry style
        self.style.configure("TEntry",
            fieldbackground=theme['entry_bg'],
            foreground=theme['entry_fg'],
            selectbackground=theme['select_bg'],
            selectforeground=theme['select_fg']
        )
        
        # Configure Combobox style
        self.style.configure("TCombobox",
            fieldbackground=theme['entry_bg'],
            foreground=theme['entry_fg'],
            selectbackground=theme['select_bg'],
            selectforeground=theme['select_fg'],
            background=theme['entry_bg']
        )
        
        # Configure Notebook and tabs
        self.style.configure("TNotebook",
            background=theme['bg']
        )
        
        self.style.configure("TNotebook.Tab",
            background=theme['tab_bg'],
            foreground="#000000",  # Always black text
            padding=[10, 5],
            font=("Helvetica", 11, "bold")
        )
        
        # Text widget colors
        text_opts = {
            'background': theme['entry_bg'],
            'foreground': theme['entry_fg'],
            'insertbackground': theme['entry_fg'],
            'selectbackground': theme['select_bg'],
            'selectforeground': theme['select_fg']
        }
        
        # Apply Text widget styling recursively
        self._apply_text_styling(self, text_opts)
        
        # Force refresh
        self.configure(bg=theme['bg'])
        self.update_idletasks()

    def _apply_text_styling(self, widget, text_opts):
        """Recursively apply text styling to all Text widgets."""
        if isinstance(widget, tk.Text):
            for key, value in text_opts.items():
                widget[key] = value
        
        # Recursively apply to all children
        for child in widget.winfo_children():
            self._apply_text_styling(child, text_opts)

    def setup_custom_styles(self):
        """Adds hover effects and ensures better contrast on widgets."""
        # Configure default tab style
        self.style.configure("TNotebook.Tab", 
                           padding=[10, 5],
                           font=("Helvetica", 11, "bold"))

        # Configure the notebook itself
        self.style.configure("TNotebook", tabmargins=[2, 5, 2, 0])

        # Hover effects for buttons
        self.style.map("TButton",
            foreground=[
                ("active", "!disabled", "#ffffff"),
                ("pressed", "#ffffff")
            ],
            background=[
                ("active", "!disabled", "#505050"),
                ("pressed", "#606060")
            ]
        )

        # Entry field focus effects
        self.style.map("TEntry",
            fieldbackground=[
                ("focus", "!disabled", "#505050"),
                ("disabled", "#2b2b2b")
            ],
            foreground=[
                ("focus", "!disabled", "#ffffff"),
                ("disabled", "#808080")
            ]
        )

        # Combobox hover and focus effects
        self.style.map("TCombobox",
            fieldbackground=[
                ("focus", "!disabled", "#505050"),
                ("disabled", "#2b2b2b"),
                ("readonly", "!disabled", "#404040")
            ],
            foreground=[
                ("focus", "!disabled", "#ffffff"),
                ("disabled", "#808080"),
                ("readonly", "!disabled", "#ffffff")
            ],
            selectbackground=[
                ("focus", "!disabled", "#606060")
            ],
            selectforeground=[
                ("focus", "!disabled", "#ffffff")
            ]
        )

    def show_about(self):
        messagebox.showinfo(
            "About Trace",
            "Trace v2.0\n\nAn investigation toolkit to capture website screenshots, "
            "generate summaries, create alert templates, and process transaction CSVs.\n"
            "Built with Python (tkinter)."
        )

    def check_for_updates(self):
        """Check for available updates."""
        try:
            update_info = self.update_manager.check_for_updates()
            if update_info['available']:
                if messagebox.askyesno("Update Available", 
                    f"Version {update_info['version']} is available.\n\n"
                    f"Release Notes:\n{update_info['release_notes']}\n\n"
                    "Would you like to update now?"):
                    self.install_update(update_info)
        except Exception as e:
            print(f"Error checking for updates: {e}")
    
    def install_update(self, update_info):
        """Install the available update."""
        try:
            # Download the update
            update_file = self.update_manager.download_update(update_info['assets'][0]['browser_download_url'])
            
            # Install the update
            if self.update_manager.install_update(update_file):
                if messagebox.askyesno("Update Complete", 
                    "Update installed successfully. Restart now?"):
                    self.update_manager.restart_application()
        except Exception as e:
            messagebox.showerror("Update Error", f"Failed to install update: {e}")

#################################
#        FRAME 1: PDF TOOL      #
#################################
class ScreenshotToPDFFrame(ttk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        # Frame Title
        title_label = ttk.Label(self, text="Convert Website Screenshots to PDF", font=("Helvetica", 14, "bold"))
        title_label.pack(pady=10, anchor="w")

        # Instructions
        instructions = (
            "Enter one or more URLs (one per line). Click 'Run' to open a headless Chrome browser, "
            "capture each page as a PDF, and append a banner with the source URL."
        )
        ttk.Label(self, text=instructions, wraplength=600).pack(pady=(0, 10), anchor="w")

        # URLs text area with carat
        self.url_text = tk.Text(self, height=8, wrap="word")
        self.url_text.pack(fill="x", padx=10, pady=5)
        
        # Configure text widget colors from parent's style with fallback values
        text_opts = {
            'background': self.controller.style.lookup('TEntry', 'fieldbackground') or '#ffffff',
            'foreground': self.controller.style.lookup('TEntry', 'foreground') or '#000000',
            'insertbackground': self.controller.style.lookup('TEntry', 'foreground') or '#000000',
            'selectbackground': self.controller.style.lookup('TEntry', 'selectbackground') or '#0078d7',
            'selectforeground': self.controller.style.lookup('TEntry', 'selectforeground') or '#ffffff',
            'insertwidth': 2  # Consistent carat width
        }
        
        for key, value in text_opts.items():
            self.url_text[key] = value

        # Independent Run button
        self.run_button = ttk.Button(self, text="Run", command=self.run)
        self.run_button.pack(pady=10)

        # Progress bar for feedback
        self.progress = ttk.Progressbar(self, mode="indeterminate")
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.pack_forget()

    def run(self):
        """Called when the user clicks the 'Run' button or from 'Run Case'."""
        raw_urls = self.url_text.get("1.0", tk.END).strip()
        urls = [line.strip() for line in raw_urls.splitlines() if line.strip()]

        if not urls:
            messagebox.showwarning("Missing URLs", "Please enter at least one URL.")
            return

        # Start the progress bar
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.start(10)

        output_dir = self.controller.get_investigation_dir()
        os.makedirs(output_dir, exist_ok=True)

        # Get the file prefix
        prefix = self.controller.get_file_prefix()

        # Set up headless Chrome (once, rather than repeatedly in the loop)
        options = Options()
        options.add_argument("--headless")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1920,1080")
        options.add_argument("--no-sandbox")

        for url in urls:
            try:
                driver = webdriver.Chrome(options=options)
                driver.get(url)
                time.sleep(2)

                # Insert a banner with the source URL
                banner = f"""
                    var banner = document.createElement('div');
                    banner.textContent = 'Source URL: {url}';
                    banner.style.padding = '10px';
                    banner.style.backgroundColor = '#f2f2f2';
                    banner.style.fontSize = '12px';
                    banner.style.fontFamily = 'monospace';
                    banner.style.borderBottom = '1px solid #ccc';
                    document.body.prepend(banner);
                """
                driver.execute_script(banner)
                time.sleep(1.5)

                pdf_data = driver.execute_cdp_cmd("Page.printToPDF", {
                    "landscape": False,
                    "displayHeaderFooter": False,
                    "printBackground": True,
                    "preferCSSPageSize": True,
                    "paperWidth": 8.27,      # A4 width
                    "paperHeight": 11.69     # A4 height
                })
                filename = f"{prefix}{url.split('//')[-1].split('/')[0]}.pdf"
                filepath = os.path.join(output_dir, filename)

                with open(filepath, 'wb') as f:
                    f.write(base64.b64decode(pdf_data['data']))

                driver.quit()

            except Exception as e:
                messagebox.showerror("Error", f"Failed to convert:\n{url}\n\n{e}")

        self.progress.stop()
        self.progress.pack_forget()
        messagebox.showinfo("Done", f"PDFs saved to:\n{output_dir}")


#################################
#     FRAME 2: GPT SUMMARY      #
#################################
class CustomerSummaryFrame(ttk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        # Create a scrollable frame
        self.canvas = tk.Canvas(self)
        self.scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        # Pack the canvas and scrollbar
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        # Title
        label = ttk.Label(self.scrollable_frame, text="Glean Summary", font=("Helvetica", 14, "bold"))
        label.pack(pady=10, anchor="w")

        # Instructions
        instructions = (
            "As of Trace v1.0, this feature is PENDING Okta access. Enter a Customer ID to fetch an investigation-focused summary. "
            "The summary is then saved as a Word document."
        )
        ttk.Label(self.scrollable_frame, text=instructions, wraplength=600).pack(anchor="w", padx=10, pady=(0, 10))

        # Customer ID entry with carat
        self.customer_entry = ttk.Entry(self.scrollable_frame)
        self.customer_entry.pack(fill="x", padx=10, pady=5)
        self.customer_entry.focus()

        # Run button
        self.summary_button = ttk.Button(self.scrollable_frame, text="Generate Summary", command=self.run_customer_summary)
        self.summary_button.pack(pady=5)

        # Progress bar
        self.progress = ttk.Progressbar(self.scrollable_frame, mode="indeterminate")
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.pack_forget()

    def run(self):
        """Called when the tab is double-clicked for 'Run Case'."""
        customer_id = self.customer_entry.get().strip()
        if customer_id:
            self.run_customer_summary()

    def run_customer_summary(self):
        """Run the customer summary generation."""
        customer_id = self.customer_entry.get().strip()
        if not customer_id:
            messagebox.showwarning("Missing ID", "Please enter a Customer ID.")
            return

        prompt = f"Please provide a summary of customer ID {customer_id}. Tailor your response for an analyst on the Investigation Operations team."
        api_url = "https://chatai.circle.com/api/proxy/open_ai/v1/chat/completions"

        headers = {
            "Authorization": "Bearer eyJhbGciOiJkaXIiLCJlbmMiOiJBMjU2R0NNIn0..UU-SRFp1e4P0onrx.F1i9VAT5xUDDvQKDvK0RZgZexIJpmj08-VV8v3GoWi5MvZSHHmiwBC_3SQftEHbfSeaXXZjTwSpCqzbJkKcGxibxye-VeDFHMmniU6-yNOgcIsbQOZ3yhOPUT5u5O96CSfxICTvNFs4DbjxJa228saTJDsgDTYBX6_6FHnY7Rr5LxVn6ZLsj3ExHbVGeYnhd6HZlEu-KCNjWD2tZ9TygnYjjJn-GT5gLDlbuA4kg9zgKyld9i554tm8.jcAJKm_ocrFyRl95GdnnXg",
            "Content-Type": "application/json"
        }

        payload = {
            "model": "gpt-4o",
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.2,
            "max_tokens": 800
        }

        # Show progress bar
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.start(10)
        self.update_idletasks()

        try:
            response = requests.post(api_url, headers=headers, json=payload)
            response.raise_for_status()
            data = response.json()
            summary = data["choices"][0]["message"]["content"]

            # Get the file prefix
            prefix = self.controller.get_file_prefix()

            # Save to Word doc
            doc = Document()
            doc.add_heading(f"Customer Summary: {customer_id}", level=1)
            doc.add_paragraph(summary)

            output_folder = self.controller.get_investigation_dir()
            os.makedirs(output_folder, exist_ok=True)
            file_path = os.path.join(output_folder, f"{prefix}{customer_id}_Summary.docx")
            doc.save(file_path)

            messagebox.showinfo("Success", f"Summary saved to:\n{file_path}")

        except requests.exceptions.HTTPError as e:
            messagebox.showerror("API Error", f"Failed to fetch summary:\n{e}\n\n{response.status_code} - {response.text}")
        except Exception as ex:
            messagebox.showerror("Error", f"An unexpected error occurred:\n{ex}")
        finally:
            self.progress.stop()
            self.progress.pack_forget()

#################################
#     FRAME 3: ALERT DOCS       #
#################################
class AlertFrame(ttk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        
        # Create notebook for sub-tabs
        self.notebook = ttk.Notebook(self)
        self.notebook.pack(fill="both", expand=True)
        
        # Create alert generation frame (now first)
        self.alert_gen_frame = AlertGenerationFrame(self.notebook, controller)
        self.notebook.add(self.alert_gen_frame, text="Generate Alert")
        
        # Create template management frame (now second)
        self.template_frame = TemplateManagementFrame(self.notebook, controller)
        self.notebook.add(self.template_frame, text="Template Management")
    
    def run(self):
        """Called when the tab is double-clicked for 'Run Case'."""
        self.alert_gen_frame.generate_alert()

class TemplateManagementFrame(ttk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "alert_templates")
        os.makedirs(self.template_dir, exist_ok=True)
        
        # Convert existing .docx templates to .txt if needed
        self.convert_docx_templates()
        
        # Create main container with padding
        main_container = ttk.Frame(self, padding="10")
        main_container.pack(fill="both", expand=True)
        
        # Left panel for template list and controls
        left_panel = ttk.Frame(main_container)
        left_panel.pack(side="left", fill="y", padx=(0, 10))
        
        # Template list
        ttk.Label(left_panel, text="Alert Templates", font=("Helvetica", 12, "bold")).pack(anchor="w", pady=(0, 5))
        self.template_listbox = tk.Listbox(left_panel, width=30, height=15)
        self.template_listbox.pack(fill="y", expand=True)
        self.template_listbox.bind('<<ListboxSelect>>', self.on_template_select)
        
        # Template controls
        controls_frame = ttk.Frame(left_panel)
        controls_frame.pack(fill="x", pady=5)
        
        ttk.Button(controls_frame, text="New", command=self.new_template).pack(side="left", padx=2)
        ttk.Button(controls_frame, text="Save", command=self.save_template).pack(side="left", padx=2)
        ttk.Button(controls_frame, text="Delete", command=self.delete_template).pack(side="left", padx=2)
        
        # Right panel for template content and placeholders
        right_panel = ttk.Frame(main_container)
        right_panel.pack(side="right", fill="both", expand=True)
        
        # Template editor
        editor_frame = ttk.LabelFrame(right_panel, text="Template Editor", padding="5")
        editor_frame.pack(fill="both", expand=True, pady=(0, 10))
        
        # Template name
        name_frame = ttk.Frame(editor_frame)
        name_frame.pack(fill="x", pady=5)
        ttk.Label(name_frame, text="Template Name:").pack(side="left")
        self.template_name = ttk.Entry(name_frame)
        self.template_name.pack(side="left", fill="x", expand=True, padx=5)
        self.template_name.bind('<FocusIn>', lambda e: self.template_name.configure(insertwidth=2))
        self.template_name.bind('<FocusOut>', lambda e: self.template_name.configure(insertwidth=0))
        
        # Template content
        content_frame = ttk.Frame(editor_frame)
        content_frame.pack(fill="both", expand=True, pady=5)
        ttk.Label(content_frame, text="Template Content:").pack(anchor="w")
        self.template_content = tk.Text(content_frame, wrap="word", height=10)
        self.template_content.pack(fill="both", expand=True)
        
        # Configure text widget colors from parent's style
        text_opts = {
            'background': self.controller.style.lookup('TEntry', 'fieldbackground'),
            'foreground': self.controller.style.lookup('TEntry', 'foreground'),
            'insertbackground': self.controller.style.lookup('TEntry', 'foreground'),
            'selectbackground': self.controller.style.lookup('TEntry', 'selectbackground'),
            'selectforeground': self.controller.style.lookup('TEntry', 'selectforeground'),
            'insertwidth': 2  # Consistent carat width
        }
        
        for key, value in text_opts.items():
            self.template_content[key] = value
            
        # Preview text widget with same styling
        preview_frame = ttk.LabelFrame(right_panel, text="Preview", padding="5")
        preview_frame.pack(fill="both", expand=True, pady=(0, 10))
        
        self.preview_text = tk.Text(preview_frame, wrap="word", height=10, state="disabled")
        self.preview_text.pack(fill="both", expand=True)
        
        # Apply same styling to preview text
        for key, value in text_opts.items():
            self.preview_text[key] = value
        
        # Placeholder Key section
        key_frame = ttk.LabelFrame(right_panel, text="Available Placeholders", padding="5")
        key_frame.pack(fill="x", pady=(0, 10))
        
        # Create a cleaner grid of placeholders
        placeholders = [
            ("{customer_id}", "Customer ID"),
            ("{customer_name}", "Customer Name"),
            ("{start_date}", "Start Date"),
            ("{end_date}", "End Date"),
            ("{account_purpose}", "Account Purpose")
        ]
        
        # Create a frame for the grid with better spacing
        grid_frame = ttk.Frame(key_frame)
        grid_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Add a header with better styling
        header_frame = ttk.Frame(grid_frame)
        header_frame.pack(fill="x", pady=(0, 10))
        ttk.Label(header_frame, text="Select and copy placeholders:", 
                 font=("Helvetica", 10, "bold")).pack(side="left")
        
        # Create a two-row, two-column grid
        for i in range(0, len(placeholders), 2):
            row_frame = ttk.Frame(grid_frame)
            row_frame.pack(fill="x", pady=2)
            
            # First column
            col1_frame = ttk.Frame(row_frame)
            col1_frame.pack(side="left", expand=True, fill="x")
            
            # Create a read-only text widget for the first placeholder with context menu
            placeholder1 = tk.Text(col1_frame, height=1, width=15, font=("Courier", 10, "bold"),
                                 wrap="none", relief="flat", padx=2, pady=2)
            placeholder1.insert("1.0", placeholders[i][0])
            placeholder1.config(state="disabled")  # Make it read-only
            
            # Add context menu for placeholder1
            context_menu1 = tk.Menu(placeholder1, tearoff=0)
            context_menu1.add_command(label="Copy", command=lambda: self.copy_selected_text(placeholder1))
            placeholder1.bind("<Button-3>", lambda e: self.show_context_menu(e, context_menu1))
            
            placeholder1.pack(side="left")
            
            ttk.Label(col1_frame, text=" = ").pack(side="left")
            ttk.Label(col1_frame, text=placeholders[i][1]).pack(side="left")
            
            # Second column (if exists)
            if i + 1 < len(placeholders):
                col2_frame = ttk.Frame(row_frame)
                col2_frame.pack(side="left", expand=True, fill="x", padx=(20, 0))
                
                # Create a read-only text widget for the second placeholder with context menu
                placeholder2 = tk.Text(col2_frame, height=1, width=15, font=("Courier", 10, "bold"),
                                     wrap="none", relief="flat", padx=2, pady=2)
                placeholder2.insert("1.0", placeholders[i+1][0])
                placeholder2.config(state="disabled")  # Make it read-only
                
                # Add context menu for placeholder2
                context_menu2 = tk.Menu(placeholder2, tearoff=0)
                context_menu2.add_command(label="Copy", command=lambda: self.copy_selected_text(placeholder2))
                placeholder2.bind("<Button-3>", lambda e: self.show_context_menu(e, context_menu2))
                
                placeholder2.pack(side="left")
                
                ttk.Label(col2_frame, text=" = ").pack(side="left")
                ttk.Label(col2_frame, text=placeholders[i+1][1]).pack(side="left")
        
        # Load existing templates
        self.load_templates()
        
        # Bind template content changes to preview
        self.template_content.bind('<<Modified>>', self.update_preview)
    
    def convert_docx_templates(self):
        """Convert existing .docx templates to .txt format."""
        for filename in os.listdir(self.template_dir):
            if filename.endswith('.docx'):
                docx_path = os.path.join(self.template_dir, filename)
                txt_path = os.path.join(self.template_dir, filename[:-5] + '.txt')
                
                # Only convert if .txt doesn't exist
                if not os.path.exists(txt_path):
                    try:
                        doc = Document(docx_path)
                        content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                        with open(txt_path, 'w') as f:
                            f.write(content)
                    except Exception as e:
                        print(f"Error converting {filename}: {e}")
    
    def load_templates(self):
        """Load existing templates into the listbox."""
        self.template_listbox.delete(0, tk.END)
        templates = []
        
        # Get all .txt files
        for template in os.listdir(self.template_dir):
            if template.endswith('.txt'):
                # Remove the "template_#_" prefix if it exists
                display_name = template[:-4]  # Remove .txt
                if display_name.startswith('template_'):
                    # Split by underscore and take everything after the second underscore
                    parts = display_name.split('_', 2)
                    if len(parts) > 2:
                        display_name = parts[2]
                templates.append((display_name, template[:-4]))  # Store both display name and full name
        
        # Sort templates by number if they start with a number
        def get_template_number(name_tuple):
            try:
                return int(name_tuple[1].split('_')[1])  # Use the full name for sorting
            except (ValueError, IndexError):
                return float('inf')
        
        templates.sort(key=get_template_number)
        
        for display_name, _ in templates:
            self.template_listbox.insert(tk.END, display_name)
        
        # Store the mapping of display names to full names
        self.template_name_map = {display: full for display, full in templates}
    
    def on_template_select(self, event):
        """Handle template selection."""
        selection = self.template_listbox.curselection()
        if selection:
            display_name = self.template_listbox.get(selection[0])
            full_name = self.template_name_map.get(display_name, display_name)
            self.load_template(full_name)
    
    def load_template(self, template_name):
        """Load a template into the editor."""
        template_path = os.path.join(self.template_dir, f"{template_name}.txt")
        if os.path.exists(template_path):
            with open(template_path, 'r') as f:
                content = f.read()
            self.template_name.delete(0, tk.END)
            # Use the display name (without prefix) in the editor
            display_name = template_name
            if display_name.startswith('template_'):
                parts = display_name.split('_', 2)
                if len(parts) > 2:
                    display_name = parts[2]
            self.template_name.insert(0, display_name)
            self.template_content.delete('1.0', tk.END)
            self.template_content.insert('1.0', content)
            self.update_preview()
    
    def new_template(self):
        """Create a new template."""
        self.template_name.delete(0, tk.END)
        self.template_content.delete('1.0', tk.END)
        self.template_name.focus()
    
    def delete_template(self):
        """Delete the selected template."""
        selection = self.template_listbox.curselection()
        if selection:
            template_name = self.template_listbox.get(selection[0])
            if messagebox.askyesno("Confirm Delete", f"Are you sure you want to delete '{template_name}'?"):
                template_path = os.path.join(self.template_dir, f"{template_name}.txt")
                if os.path.exists(template_path):
                    os.remove(template_path)
                    self.load_templates()
                    self.new_template()
    
    def save_template(self):
        """Save the current template."""
        name = self.template_name.get().strip()
        content = self.template_content.get('1.0', tk.END).strip()
        
        if not name:
            messagebox.showwarning("Missing Name", "Please enter a template name.")
            return
        
        if not content:
            messagebox.showwarning("Missing Content", "Please enter template content.")
            return
        
        # Check if we're editing an existing template
        selection = self.template_listbox.curselection()
        if selection:
            # We're editing an existing template
            display_name = self.template_listbox.get(selection[0])
            full_name = self.template_name_map.get(display_name, display_name)
            
            # If the name hasn't changed, use the existing full name
            if name == display_name:
                template_path = os.path.join(self.template_dir, f"{full_name}.txt")
            else:
                # Name has changed, create new template with new name
                if not name.startswith('template_'):
                    # Find the next available number
                    existing_numbers = []
                    for template in os.listdir(self.template_dir):
                        if template.startswith('template_') and template.endswith('.txt'):
                            try:
                                num = int(template.split('_')[1])
                                existing_numbers.append(num)
                            except (ValueError, IndexError):
                                continue
                    
                    next_num = 1
                    while next_num in existing_numbers:
                        next_num += 1
                    
                    name = f"template_{next_num:02d}_{name}"
                
                template_path = os.path.join(self.template_dir, f"{name}.txt")
                
                # Remove old template file
                old_path = os.path.join(self.template_dir, f"{full_name}.txt")
                if os.path.exists(old_path):
                    os.remove(old_path)
        else:
            # Creating a new template
            if not name.startswith('template_'):
                # Find the next available number
                existing_numbers = []
                for template in os.listdir(self.template_dir):
                    if template.startswith('template_') and template.endswith('.txt'):
                        try:
                            num = int(template.split('_')[1])
                            existing_numbers.append(num)
                        except (ValueError, IndexError):
                            continue
                
                next_num = 1
                while next_num in existing_numbers:
                    next_num += 1
                
                name = f"template_{next_num:02d}_{name}"
            
            template_path = os.path.join(self.template_dir, f"{name}.txt")
        
        # Save the template
        with open(template_path, 'w') as f:
            f.write(content)
        
        # Refresh the template list
        self.load_templates()
        
        # Select the newly saved template
        display_name = name
        if display_name.startswith('template_'):
            parts = display_name.split('_', 2)
            if len(parts) > 2:
                display_name = parts[2]
        
        # Find and select the template in the listbox
        for i in range(self.template_listbox.size()):
            if self.template_listbox.get(i) == display_name:
                self.template_listbox.selection_clear(0, tk.END)
                self.template_listbox.selection_set(i)
                self.template_listbox.see(i)
                break
        
        messagebox.showinfo("Success", f"Template '{display_name}' saved successfully.")
    
    def update_preview(self, event=None):
        """Update the preview with current template content."""
        content = self.template_content.get('1.0', tk.END)
        self.preview_text.config(state="normal")
        self.preview_text.delete('1.0', tk.END)
        self.preview_text.insert('1.0', content)
        self.preview_text.config(state="disabled")
        if event:
            self.template_content.edit_modified(False)

    def show_context_menu(self, event, menu):
        """Show the context menu at the mouse position."""
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()
    
    def copy_selected_text(self, text_widget):
        """Copy the selected text to clipboard."""
        try:
            selected_text = text_widget.get("sel.first", "sel.last")
            self.clipboard_clear()
            self.clipboard_append(selected_text)
        except tk.TclError:
            # If no text is selected, copy the entire content
            text_widget.config(state="normal")
            content = text_widget.get("1.0", "end-1c")
            text_widget.config(state="disabled")
            self.clipboard_clear()
            self.clipboard_append(content)

class AlertGenerationFrame(ttk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller
        self.template_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "alert_templates")
        
        # Create main container with padding
        main_container = ttk.Frame(self, padding="10")
        main_container.pack(fill="both", expand=True)
        
        # Template selection
        template_frame = ttk.LabelFrame(main_container, text="Select Template", padding="5")
        template_frame.pack(fill="x", pady=(0, 10))
        
        # Create a dropdown for template selection
        self.template_var = tk.StringVar()
        self.template_dropdown = ttk.Combobox(template_frame, textvariable=self.template_var, state="readonly")
        self.template_dropdown.pack(fill="x", pady=5)
        self.load_templates()
        
        # Alert information
        info_frame = ttk.LabelFrame(main_container, text="Alert Information", padding="5")
        info_frame.pack(fill="both", expand=True)
        
        # Entry fields
        self.alert_customer_id = self.labeled_entry(info_frame, "Customer ID")
        self.alert_customer_name = self.labeled_entry(info_frame, "Customer Name")
        self.alert_start_date = self.labeled_entry(info_frame, "Start Date (MM/DD/YYYY)")
        self.alert_end_date = self.labeled_entry(info_frame, "End Date (MM/DD/YYYY)")
        self.alert_account_purpose = self.labeled_entry(info_frame, "Account Purpose")
        
        # Generate button
        ttk.Button(main_container, text="Generate Alert", command=self.generate_alert).pack(pady=10)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_container, mode="indeterminate")
        self.progress.pack(fill="x", pady=5)
        self.progress.pack_forget()
    
    def labeled_entry(self, parent, label_text):
        """Create a labeled entry field."""
        frame = ttk.Frame(parent)
        frame.pack(fill="x", pady=2)
        ttk.Label(frame, text=label_text).pack(side="left")
        entry = ttk.Entry(frame)
        entry.pack(side="left", fill="x", expand=True, padx=5)
        return entry
    
    def load_templates(self):
        """Load existing templates into the dropdown with clean display names."""
        templates = []
        
        # Get all .txt files
        for template in os.listdir(self.template_dir):
            if template.endswith('.txt'):
                # Remove the "template_#_" prefix if it exists
                display_name = template[:-4]  # Remove .txt
                if display_name.startswith('template_'):
                    # Split by underscore and take everything after the second underscore
                    parts = display_name.split('_', 2)
                    if len(parts) > 2:
                        display_name = parts[2]
                templates.append((display_name, template[:-4]))  # Store both display name and full name
        
        # Sort templates by number if they start with a number
        def get_template_number(name_tuple):
            try:
                return int(name_tuple[1].split('_')[1])  # Use the full name for sorting
            except (ValueError, IndexError):
                return float('inf')
        
        templates.sort(key=get_template_number)
        
        # Update dropdown values with clean display names
        self.template_dropdown['values'] = [display_name for display_name, _ in templates]
        
        # Store the mapping of display names to full names
        self.template_name_map = {display: full for display, full in templates}
        
        # Configure the combobox style to match the theme
        self.template_dropdown.configure(
            background=self.controller.style.lookup('TEntry', 'fieldbackground'),
            foreground=self.controller.style.lookup('TEntry', 'foreground')
        )
    
    def generate_alert(self):
        """Generate an alert document using the selected template."""
        template_name = self.template_var.get()
        if not template_name:
            messagebox.showwarning("Missing Template", "Please select a template.")
            return
        
        full_name = self.template_name_map.get(template_name, template_name)
        template_path = os.path.join(self.template_dir, f"{full_name}.txt")
        
        if not os.path.exists(template_path):
            messagebox.showwarning("Missing Template", "Selected template not found.")
            return
        
        context = {
            "customer_id": self.alert_customer_id.get().strip(),
            "customer_name": self.alert_customer_name.get().strip(),
            "start_date": self.alert_start_date.get().strip(),
            "end_date": self.alert_end_date.get().strip(),
            "account_purpose": self.alert_account_purpose.get().strip()
        }
        
        if not all(context.values()):
            messagebox.showwarning("Missing Info", "Please complete all fields.")
            return
        
        # Show progress
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.start(10)
        self.update_idletasks()
        
        try:
            # Load and fill the template
            doc = Document()
            
            # Add template content
            with open(template_path, 'r') as f:
                template_content = f.read()
            
            # Replace placeholders in template
            for key, value in context.items():
                template_content = template_content.replace(f"{{{key}}}", value)
            
            # Add template content to document
            doc.add_paragraph(template_content)
            
            # Save final doc with CustomerID_Name prefix
            output_folder = self.controller.get_investigation_dir()
            os.makedirs(output_folder, exist_ok=True)
            
            # Get the CustomerID_Name prefix
            prefix = self.controller.get_file_prefix()
            
            # Create filename using just the CustomerID_Name and "Narrative"
            filename = f"{prefix}Narrative.docx"
            filepath = os.path.join(output_folder, filename)
            doc.save(filepath)
            
            messagebox.showinfo("Success", f"Alert template saved to:\n{filepath}")
            
        except Exception as e:
            messagebox.showerror("Generation Error", f"Failed to generate alert:\n{e}")
        finally:
            self.progress.stop()
            self.progress.pack_forget()

#################################
#   FRAME 4: TRANSACTION CSV    #
#################################
class TransactionCSVFrame(ttk.Frame):
    def __init__(self, parent, controller):
        super().__init__(parent)
        self.controller = controller

        # To store selected CSV path
        self.csv_path = None

        # Title
        ttk.Label(self, text="Process Transaction CSV", font=("Helvetica", 14, "bold")).pack(anchor="w", pady=10)

        # Instruction
        instructions = (
            "1) Click 'Upload CSV' to select your transactions file.\n"
            "2) Click 'Run' to process the CSV: filter out failed/denied, categorize, and create a summary breakdown."
        )
        ttk.Label(self, text=instructions, wraplength=700).pack(anchor="w", padx=10, pady=(0, 10))

        # Buttons
        upload_button = ttk.Button(self, text="Upload CSV", command=self.choose_csv_file)
        upload_button.pack(pady=5)

        # Label to display selected file
        self.file_label = ttk.Label(self, text="No file selected", font=("Helvetica", 10))
        self.file_label.pack(pady=5)

        self.run_button = ttk.Button(self, text="Run", command=self.run)
        self.run_button.pack(pady=10)

        # Progress bar
        self.progress = ttk.Progressbar(self, mode="indeterminate")
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.pack_forget()

    def choose_csv_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if file_path:
            self.csv_path = file_path
            # Display just the filename, not the full path
            filename = os.path.basename(file_path)
            self.file_label.config(text=f"Selected: {filename}")
            messagebox.showinfo("CSV Selected", f"Selected CSV:\n{file_path}")

    def run(self):
        """Called when the user clicks 'Run' or from 'Run Case' to process the CSV."""
        if not self.csv_path:
            messagebox.showwarning("No CSV Selected", "Please upload a CSV file before running.")
            return

        # Show progress
        self.progress.pack(fill="x", padx=20, pady=(0,10))
        self.progress.start(10)
        self.update_idletasks()

        try:
            df = pd.read_csv(self.csv_path)
            df.columns = df.columns.str.strip().str.lower()

            # Ensure correct types and fill values
            if 'create_date' in df.columns:
                df['create_date'] = pd.to_datetime(df['create_date'], errors='coerce')

            for col in ["activity_state", "activity", "funding_type", "destination_type",
                        "from_address", "destination", "blockchain", "to_currency"]:
                if col in df.columns:
                    df[col] = df[col].astype(str).str.lower().fillna('')

            numeric_cols = ["usd_value_updated", "original_currency_amount_updated"]
            for col in numeric_cols:
                if col in df.columns:
                    df[col] = pd.to_numeric(df[col], errors='coerce').fillna(0)

            # Filter out failed/denied
            if 'activity_state' in df.columns:
                valid_rows = df[~df['activity_state'].isin(['failed', 'denied'])].copy()
            else:
                valid_rows = df.copy()

            # Define categories
            categories = {
                'onchainReceive': lambda r: r.get('activity','') == 'receive' and r.get('funding_type','') == 'blockchain',
                'onchainSpend':   lambda r: r.get('activity','') == 'spend',
                'incomingWire':   lambda r: r.get('activity','') == 'receive' and r.get('funding_type','') == 'fiat_account',
                'outgoingWire':   lambda r: r.get('activity','') == 'eft_transfer'
            }

            valid_rows['category'] = 'other'
            for cat, func in categories.items():
                mask = valid_rows.apply(func, axis=1)
                valid_rows.loc[mask, 'category'] = cat

            # Add month-year and amount columns
            if 'create_date' in valid_rows.columns:
                valid_rows['month_year'] = valid_rows['create_date'].dt.to_period("M").astype(str)

            if 'usd_value_updated' in valid_rows.columns:
                valid_rows['amount'] = valid_rows['usd_value_updated']
            else:
                valid_rows['amount'] = 0

            if 'original_currency_amount_updated' in valid_rows.columns:
                zero_mask = valid_rows['amount'] == 0
                valid_rows.loc[zero_mask, 'amount'] = valid_rows['original_currency_amount_updated']

            # Generate summary
            def format_currency(val):
                return "${:,.2f}".format(val) if pd.notnull(val) else "$0.00"

            summary_lines = []
            for cat in categories.keys():
                cat_rows = valid_rows[valid_rows['category'] == cat]
                if not cat_rows.empty:
                    total = cat_rows['amount'].sum()
                    count = len(cat_rows)
                    min_amt = cat_rows['amount'].min()
                    max_amt = cat_rows['amount'].max()
                    dates = cat_rows['create_date'].dropna() if 'create_date' in cat_rows.columns else []
                    if len(dates) > 0:
                        line = (f"{count} {cat} transactions totaling {format_currency(total)}, "
                                f"from {dates.min().date()} to {dates.max().date()}.")
                    else:
                        line = f"{count} {cat} transactions totaling {format_currency(total)}."
                    summary_lines.append(line)

            # Additional breakdown by 'activity'
            if 'activity' in valid_rows.columns:
                for act in valid_rows['activity'].unique():
                    act_rows = valid_rows[valid_rows['activity'] == act]
                    if not act_rows.empty:
                        total = act_rows['amount'].sum()
                        count = len(act_rows)
                        line = f"{count} '{act}' transactions totaling {format_currency(total)}."
                        summary_lines.append(line)

            # Destination stats
            if 'destination_type' in valid_rows.columns:
                dest_blockchain = valid_rows[valid_rows['destination_type'] == 'blockchain']
                if not dest_blockchain.empty and 'destination' in dest_blockchain.columns:
                    dest_stats = dest_blockchain.groupby('destination', dropna=False).agg(
                        Count=('destination', 'size'),
                        Total_USD=('amount', 'sum')
                    ).sort_values(by='Total_USD', ascending=False).head(50)
                else:
                    dest_stats = pd.DataFrame()
            else:
                dest_stats = pd.DataFrame()

            # From address stats
            if 'from_address' in valid_rows.columns:
                from_stats = valid_rows[valid_rows['from_address'].str.lower() != 'null']
                if not from_stats.empty:
                    from_stats = from_stats.groupby('from_address', dropna=False).agg(
                        Count=('from_address', 'size'),
                        Total_USD=('amount', 'sum')
                    ).sort_values(by='Total_USD', ascending=False).head(50)
                else:
                    from_stats = pd.DataFrame()
            else:
                from_stats = pd.DataFrame()

            # Month-Year breakdown
            if 'month_year' in valid_rows.columns and 'category' in valid_rows.columns:
                month_cat = valid_rows.groupby(['month_year', 'category']).agg(
                    Count=('amount', 'size'),
                    Total_USD=('amount', 'sum'),
                    Min_USD=('amount', 'min'),
                    Max_USD=('amount', 'max')
                ).reset_index()
            else:
                month_cat = pd.DataFrame()

            # Blockchain by category
            if 'blockchain' in valid_rows.columns and 'category' in valid_rows.columns:
                chain_cat = valid_rows.groupby(['blockchain', 'category']).agg(
                    Count=('amount', 'size'),
                    Total_USD=('amount', 'sum')
                ).reset_index()
            else:
                chain_cat = pd.DataFrame()

            # Get the file prefix
            prefix = self.controller.get_file_prefix()

            # Save results
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_folder = self.controller.get_investigation_dir()
            os.makedirs(output_folder, exist_ok=True)
            output_file = os.path.join(output_folder, f"{prefix}Transactions.xlsx")

            # Strip timezone info from all datetime columns to avoid Excel writer issues
            for col in df.columns:
                if pd.api.types.is_datetime64_any_dtype(df[col]):
                    df[col] = pd.to_datetime(df[col], errors="coerce").dt.tz_localize(None)

            with pd.ExcelWriter(output_file, engine='xlsxwriter') as writer:
                pd.DataFrame({"Summary": summary_lines}).to_excel(writer, sheet_name='Summary', index=False)
                if not dest_stats.empty:
                    dest_stats.to_excel(writer, sheet_name='Destination Blockchain', index=True)
                if not from_stats.empty:
                    from_stats.to_excel(writer, sheet_name='From Address', index=True)
                if not month_cat.empty:
                    month_cat.to_excel(writer, sheet_name='Month-Year Breakdown', index=False)
                if not chain_cat.empty:
                    chain_cat.to_excel(writer, sheet_name='Blockchain by Type', index=False)
                df.to_excel(writer, sheet_name='Raw Transactions', index=False)

            messagebox.showinfo("Success", f"Summary saved to:\n{output_file}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to process transactions:\n{e}")
        finally:
            self.progress.stop()
            self.progress.pack_forget()


#################################
#           MAIN LOOP           #
#################################
if __name__ == "__main__":
    trace_app = TraceApp()
    sys.exit(trace_app.mainloop())

def check_for_updates():
    update_info = update_manager.check_for_updates()
    if update_info['available']:
        # Notify user and handle update
        pass
